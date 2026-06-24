"""Azure AI Vision Read OCR with a never-raises contract.

The public entrypoint :func:`extract_pages` turns raw document bytes into an
:class:`~di.models.OcrResult`. It is designed to *never* raise: any missing credential,
absent optional dependency, or runtime error degrades gracefully to a deterministic
fallback (a ``pypdf`` text-layer extraction when available, otherwise an empty result).

Engine resolution order:

1. ``azure-vision-read`` — Azure AI Vision Read OCR, used when
   :attr:`di.config.Settings.has_azure_vision` is true and the SDK imports.
   Azure Vision Read is image-first; for PDFs the pages are rasterized server-side or
   processed via the asynchronous Read operation (Computer Vision Read 3.2). Here we drive
   the synchronous ``ImageAnalysisClient`` (``azure-ai-vision-imageanalysis``) which returns
   a single logical page of line geometry; multipage PDFs would be handled by the async Read
   op upstream. We map whatever the SDK returns to ``OcrLine`` rows with bounding boxes.
2. ``pypdf`` — lazy text-layer extraction from a native PDF (no OCR, but free and offline).
3. ``none`` — empty result; nothing could be extracted.

All heavy/optional dependencies (azure SDK, pypdf) are imported lazily *inside* functions so
this module imports cleanly in any environment.
"""
from __future__ import annotations

import logging

from di.config import get_settings
from di.models import BBox, OcrLine, OcrResult

logger = logging.getLogger(__name__)

# Engine identifiers (kept as constants so callers can compare without typos).
ENGINE_AZURE = "azure-vision-read"
ENGINE_PYPDF = "pypdf"
ENGINE_DOCX = "docx"
ENGINE_TESSERACT = "tesseract"
ENGINE_TEXT = "text"
ENGINE_NONE = "none"


def _safe(fn):
    """Call ``fn`` returning its result, or ``None`` on any exception (OCR must never raise)."""
    try:
        return fn()
    except Exception:  # noqa: BLE001 - degrade gracefully
        logger.exception("OCR call failed; degrading")
        return None


def extract_pages(
    content: bytes,
    *,
    filename: str = "",
    mime: str | None = None,
) -> OcrResult:
    """Extract text + line geometry from ``content``; never raises.

    Args:
        content: Raw document bytes (image or PDF).
        filename: Optional original filename (informational; aids logging/heuristics).
        mime: Optional MIME type hint (e.g. ``"application/pdf"``).

    Returns:
        An :class:`~di.models.OcrResult`. ``engine`` is one of ``"azure-vision-read"``,
        ``"pypdf"``, or ``"none"``. On any failure or missing capability the function falls
        back to ``pypdf`` (if available and the bytes look like a PDF) and finally to an
        empty ``"none"`` result.
    """
    if not content:
        return _empty()

    settings = get_settings()
    azure = settings.has_azure_vision
    kind = _detect_kind(content, filename, mime)

    if kind == "image":
        if azure:
            r = _safe(lambda: _azure_read(content, filename=filename, mime=mime))
            if r is not None:
                return r
        return _image_ocr(content) or _empty()
    if kind == "pdf":
        digital = _pypdf_text_layer(content)  # selectable text layer
        if digital is not None:
            return digital
        # scanned PDF: rasterize, OCR each page with Azure (if configured) else Tesseract
        page_ocr = (lambda b: _safe(lambda: _azure_read(b))) if azure else _image_ocr
        return _pdf_ocr(content, page_ocr=page_ocr) or _empty()
    if kind == "docx":
        return _docx_extract(content) or _empty()
    if kind == "text":
        return _text_passthrough(content, mime=mime) or _empty()
    # unknown: try text, then Azure (if configured), then Tesseract
    text_result = _text_passthrough(content, mime=mime)
    if text_result is not None:
        return text_result
    if azure:
        r = _safe(lambda: _azure_read(content))
        if r is not None:
            return r
    return _image_ocr(content) or _empty()


# ---------------------------------------------------------------------------
# Azure OCR — Computer Vision Read v3.2 REST API (no SDK; plain httpx).
# Speaks the v3.2 contract, so it works unchanged against real Azure
# (https://<resource>.cognitiveservices.azure.com/) or the local mock container.
# ---------------------------------------------------------------------------
def _azure_read(content: bytes, *, filename: str = "", mime: str | None = None) -> OcrResult | None:
    """Call the Azure Computer Vision **Read v3.2** REST API and map the result.

    Async pattern: POST bytes to ``/vision/v3.2/read/analyze`` (→ 202 + ``Operation-Location``),
    then poll ``GET <Operation-Location>`` until ``status`` is ``succeeded``/``failed``. Uses
    ``httpx`` only — no Azure SDK. Returns ``None`` on any error so the caller can fall back.
    """
    import time

    import httpx

    settings = get_settings()
    base = settings.azure_vision_endpoint.rstrip("/")
    key = settings.azure_vision_key
    analyze_url = f"{base}/vision/v3.2/read/analyze"
    auth = {"Ocp-Apim-Subscription-Key": key}
    try:
        with httpx.Client(timeout=60.0) as client:
            resp = client.post(
                analyze_url,
                headers={**auth, "Content-Type": "application/octet-stream"},
                content=content,
            )
            resp.raise_for_status()
            op_location = (resp.headers.get("Operation-Location")
                           or resp.headers.get("operation-location"))
            if not op_location:
                logger.warning("Azure Read: missing Operation-Location header")
                return None
            for _ in range(120):
                poll = client.get(op_location, headers=auth)
                poll.raise_for_status()
                data = poll.json()
                status = (data.get("status") or "").lower()
                if status == "succeeded":
                    return _map_v32(data)
                if status == "failed":
                    logger.warning("Azure Read reported status=failed")
                    return None
                time.sleep(0.5)
    except httpx.HTTPError:
        logger.exception("Azure OCR v3.2 request failed")
        return None
    logger.warning("Azure Read polling timed out")
    return None


def _bbox_from_polygon(poly: object, page: int) -> BBox | None:
    """v3.2 ``boundingBox`` is a flat [x1,y1,x2,y2,x3,y3,x4,y4]; collapse to an axis-aligned BBox."""
    if not isinstance(poly, (list, tuple)) or len(poly) < 8:
        return None
    try:
        xs = [float(v) for v in poly[0::2]]
        ys = [float(v) for v in poly[1::2]]
    except (TypeError, ValueError):
        return None
    return BBox(page=page, x0=min(xs), y0=min(ys), x1=max(xs), y1=max(ys))


def _map_v32(data: dict) -> OcrResult | None:
    """Map an Azure Read v3.2 ``analyzeResults`` payload to :class:`OcrResult`."""
    read_results = (data.get("analyzeResult") or {}).get("readResults") or []
    lines: list[OcrLine] = []
    texts: list[str] = []
    for rr in read_results:
        page = int(rr.get("page", 1) or 1)
        for ln in rr.get("lines") or []:
            text = ln.get("text", "") or ""
            texts.append(text)
            words = ln.get("words") or []
            confs = [w["confidence"] for w in words
                     if isinstance(w.get("confidence"), (int, float))]
            conf = (sum(confs) / len(confs)) if confs else None
            lines.append(OcrLine(text=text, page=page,
                                 bbox=_bbox_from_polygon(ln.get("boundingBox"), page),
                                 confidence=conf))
    if not lines:
        return None
    return OcrResult(engine=ENGINE_AZURE, pages=len(read_results) or 1,
                     text="\n".join(texts), lines=lines)


# ---------------------------------------------------------------------------
# Deterministic fallback
# ---------------------------------------------------------------------------
def _empty() -> OcrResult:
    return OcrResult(engine=ENGINE_NONE, pages=0, text="", lines=[])


def _detect_kind(content: bytes, filename: str, mime: str | None) -> str:
    """Classify the upload into pdf | docx | image | text | unknown (magic / mime / extension)."""
    fn = filename.lower()
    m = (mime or "").lower()
    if content[:5] == b"%PDF-" or "pdf" in m or fn.endswith(".pdf"):
        return "pdf"
    if fn.endswith(".docx") or "word" in m or "officedocument.wordprocessing" in m:
        return "docx"
    if (content[:8].startswith(b"\x89PNG") or content[:3] == b"\xff\xd8\xff"
            or m.startswith("image/")
            or fn.endswith((".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp", ".heic"))):
        return "image"
    if m.startswith("text") or _looks_like_text(content, mime) is not None:
        return "text"
    return "unknown"


def _docx_extract(content: bytes) -> OcrResult | None:
    """Extract text from a .docx via python-docx (paragraphs + table cells). Lazy + guarded."""
    try:
        import docx  # python-docx
    except ImportError:
        return None
    import io
    try:
        document = docx.Document(io.BytesIO(content))
    except Exception:  # noqa: BLE001 - not a valid docx / parse error
        return None
    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for tbl in document.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("  ".join(cells))
    if not parts:
        return None
    lines = [OcrLine(text=p, page=1) for p in parts]
    return OcrResult(engine=ENGINE_DOCX, pages=1, text="\n".join(parts), lines=lines)


def _image_ocr(content: bytes) -> OcrResult | None:
    """OCR an image with Tesseract (pytesseract + Pillow), grouping words into lines with bboxes."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return None
    import io
    try:
        img = Image.open(io.BytesIO(content))
        data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
    except Exception:  # noqa: BLE001 - tesseract missing / unreadable image
        logger.exception("tesseract OCR failed")
        return None
    groups: dict[tuple, dict] = {}
    for i in range(len(data["text"])):
        word = (data["text"][i] or "").strip()
        if not word:
            continue
        key = (data["block_num"][i], data["par_num"][i], data["line_num"][i])
        x, y, w, hgt = data["left"][i], data["top"][i], data["width"][i], data["height"][i]
        g = groups.setdefault(key, {"w": [], "x0": x, "y0": y, "x1": x + w, "y1": y + hgt, "c": []})
        g["w"].append(word)
        g["x0"], g["y0"] = min(g["x0"], x), min(g["y0"], y)
        g["x1"], g["y1"] = max(g["x1"], x + w), max(g["y1"], y + hgt)
        try:
            g["c"].append(float(data["conf"][i]))
        except (TypeError, ValueError, KeyError):
            pass
    lines, texts = [], []
    for _, g in sorted(groups.items()):
        text = " ".join(g["w"])
        texts.append(text)
        conf = (sum(g["c"]) / len(g["c"]) / 100.0) if g["c"] else None
        lines.append(OcrLine(text=text, page=1,
                             bbox=BBox(page=1, x0=g["x0"], y0=g["y0"], x1=g["x1"], y1=g["y1"]),
                             confidence=conf))
    if not lines:
        return None
    return OcrResult(engine=ENGINE_TESSERACT, pages=1, text="\n".join(texts), lines=lines)


def _pdf_ocr(content: bytes, *, page_ocr=None) -> OcrResult | None:
    """OCR a scanned (no text-layer) PDF: rasterize via pdf2image (poppler), then OCR each page
    with ``page_ocr`` (defaults to Tesseract; Azure Vision when configured)."""
    if page_ocr is None:
        page_ocr = _image_ocr
    try:
        from pdf2image import convert_from_bytes
    except ImportError:
        return None
    try:
        images = convert_from_bytes(content, dpi=200)
    except Exception:  # noqa: BLE001 - poppler missing / render error
        logger.exception("pdf2image render failed")
        return None
    import io
    all_lines: list[OcrLine] = []
    texts: list[str] = []
    engine = ENGINE_TESSERACT
    for pageno, img in enumerate(images, start=1):
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        page = page_ocr(buf.getvalue())
        if page:
            engine = page.engine
            for ln in page.lines:
                ln.page = pageno
                all_lines.append(ln)
            texts.append(page.text)
    if not all_lines:
        return None
    return OcrResult(engine=engine, pages=len(images), text="\n\n".join(texts), lines=all_lines)


def _looks_like_text(content: bytes, mime: str | None) -> str | None:
    """Return the decoded string if the bytes are UTF-8 text (or mime is ``text/*``)."""
    if mime and mime.lower().startswith("text"):
        return content.decode("utf-8", errors="replace")
    try:
        decoded = content.decode("utf-8")
    except UnicodeDecodeError:
        return None
    if not decoded.strip():
        return None
    printable = sum(ch.isprintable() or ch.isspace() for ch in decoded)
    return decoded if printable / max(len(decoded), 1) >= 0.9 else None


def _text_passthrough(content: bytes, *, mime: str | None) -> OcrResult | None:
    """Treat already-textual input as OCR output directly.

    Covers ``text/*`` uploads and decodable UTF-8 bytes (pre-OCR'd exports, plain-text KYC
    documents, local/dev ingestion without a live OCR engine). One ``OcrLine`` per non-empty line.
    """
    decoded = _looks_like_text(content, mime)
    if decoded is None:
        return None
    lines = [OcrLine(text=ln.strip(), page=1) for ln in decoded.splitlines() if ln.strip()]
    return OcrResult(engine=ENGINE_TEXT, pages=1, text=decoded.strip(), lines=lines)


def _looks_like_pdf(content: bytes, *, filename: str, mime: str | None) -> bool:
    """Heuristic: does this look like a PDF? (magic bytes / mime / extension)."""
    if content[:5] == b"%PDF-":
        return True
    if mime and "pdf" in mime.lower():
        return True
    return filename.lower().endswith(".pdf")


def _pypdf_text_layer(content: bytes) -> OcrResult | None:
    """Lazy ``pypdf`` text-layer extraction. Returns ``None`` if pypdf is absent or fails.

    This is *not* OCR — it only recovers an embedded text layer from native PDFs. We emit one
    :class:`OcrLine` per non-empty physical line, tagged with its page, but without geometry
    (pypdf does not expose reliable per-line bounding boxes here).
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        return None

    import io

    try:
        reader = PdfReader(io.BytesIO(content))
        page_texts: list[str] = []
        lines: list[OcrLine] = []
        for page_index, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            page_texts.append(page_text)
            for raw_line in page_text.splitlines():
                stripped = raw_line.strip()
                if stripped:
                    lines.append(OcrLine(text=stripped, page=page_index, bbox=None, confidence=None))
    except Exception:  # corrupt/encrypted PDF, etc. — degrade to empty result
        logger.exception("pypdf text-layer extraction failed")
        return None

    text = "\n".join(t for t in page_texts if t).strip()
    if not text and not lines:
        return None
    return OcrResult(engine=ENGINE_PYPDF, pages=len(page_texts), text=text, lines=lines)
